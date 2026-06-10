#!/usr/bin/env python3
"""Compile the Field Guide markdown into a clean, Word-native .docx."""
import os, re, glob, io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

ROOT = "/home/user/furniture-id-app/docs/manual"
PNG  = "/tmp/dpng"
OUT  = os.path.join(ROOT, "field-guide.docx")
INK  = RGBColor(0x3A,0x2E,0x26)
GRY  = RGBColor(0x8a,0x7a,0x64)
OX   = RGBColor(0x9B,0x4B,0x3B)

SECTIONS = [
    ("PART I — THE METHOD", ["part1-method.md"]),
    ("PART II — THE FORMS · Seating", ["phase1b-sample-chapter-windsor.md","chapters/rocking-chair.md",
        "chapters/morris-chair.md","chapters/wing-chair.md","chapters/settee.md","chapters/victorian-sofa.md"]),
    ("Tables", ["chapters/parlor-table.md","chapters/drop-leaf-table.md","chapters/gateleg-table.md",
        "chapters/tilt-top-table.md","chapters/pembroke-table.md","chapters/candle-stand.md","chapters/work-table.md"]),
    ("Case and Storage", ["chapters/chest-of-drawers.md","chapters/dresser.md","chapters/dressing-table-vanity.md",
        "chapters/washstand.md","chapters/sideboard.md","chapters/china-cabinet.md","chapters/corner-cabinet.md",
        "chapters/hoosier-cabinet.md","chapters/blanket-chest.md","chapters/armoire-wardrobe.md"]),
    ("Desks", ["chapters/slant-front-desk.md","chapters/secretary-desk.md","chapters/roll-top-cylinder-desk.md"]),
    ("Beds", ["chapters/iron-bed.md","chapters/wooden-bedstead.md"]),
    ("Everyday Pieces", ["chapters/telephone-stand.md","chapters/nightstand.md"]),
    ("Clocks", ["chapters/clocks.md"]),
    ("PART III — REFERENCE ATLASES", ["atlases/fasteners.md","atlases/joinery.md","atlases/wood-species.md",
        "atlases/hardware-casters.md","atlases/finishes.md","atlases/glossary.md"]),
]
CHAPTER_PLATES = {
    "phase1b-sample-chapter-windsor.md":["windsor-back-types"],"chapters/rocking-chair.md":["rocker-types"],
    "chapters/wing-chair.md":["legs-by-period"],"chapters/victorian-sofa.md":["victorian-sofa-silhouettes"],
    "chapters/gateleg-table.md":["turned-leg-types"],"chapters/parlor-table.md":["table-base-styles"],
    "chapters/drop-leaf-table.md":["drop-leaf-supports"],"chapters/blanket-chest.md":["blanket-chest-feet"],
    "chapters/roll-top-cylinder-desk.md":["roll-top-covers"],"chapters/wooden-bedstead.md":["bedstead-subtypes"],
    "chapters/clocks.md":["clock-silhouettes"],"chapters/hoosier-cabinet.md":["hoosier-anatomy"],
    "atlases/joinery.md":["mortise-and-tenon"],"atlases/fasteners.md":["fasteners-nail-types","fasteners-screws"],
}

FILE_TREE = {
    "chapters/hoosier-cabinet.md": "hoosier-pantry-comparison",
    "chapters/sideboard.md": "sideboard-buffet-server",
}
CURRENT_FILE = ""

def read(p):
    fp=os.path.join(ROOT,p); return open(fp).read() if os.path.exists(fp) else ""

def strip_chapter(md):
    lines=md.split("\n"); i=0
    if lines and re.match(r'^#\s+(Chapter Draft:|Atlas:)', lines[0]): i=1
    while i<len(lines) and (lines[i].strip().startswith(">") or lines[i].strip()==""): i+=1
    body="\n".join(lines[i:])
    for marker in ["## Illustration list","## NOTES FOR OWNER","## Notes on format","## Illustration needs"]:
        idx=body.find(marker)
        if idx!=-1: body=body[:idx]
    return body.strip()

doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(6); sec.page_height=Inches(9)
sec.top_margin=sec.bottom_margin=Inches(0.65); sec.left_margin=sec.right_margin=Inches(0.7)
n=doc.styles['Normal']; n.font.name='Georgia'; n.font.size=Pt(11); n.font.color.rgb=INK
for i in (1,2,3):
    s=doc.styles[f'Heading {i}']; s.font.color.rgb=(OX if i==2 else INK); s.font.name='Georgia'

def add_runs(p, text):
    pat=re.compile(r'(\*\*.+?\*\*|`.+?`|\[[^\]]+?\]\([^)]+?\)|\*[^*\n]+?\*)')
    pos=0
    for m in pat.finditer(text):
        if m.start()>pos: p.add_run(text[pos:m.start()])
        tok=m.group(0)
        if tok.startswith('**'): p.add_run(tok[2:-2]).bold=True
        elif tok.startswith('`'):
            r=p.add_run(tok[1:-1]); r.font.name='Consolas'; r.font.size=Pt(10)
        elif tok.startswith('['):
            mm=re.match(r'\[(.+?)\]\((.+?)\)',tok); p.add_run(mm.group(1)).italic=True
        elif tok.startswith('*'): p.add_run(tok[1:-1]).italic=True
        pos=m.end()
    if pos<len(text): p.add_run(text[pos:])

def add_png(base):
    # owner's finished PNG in diagrams/ wins over the cairosvg rasterization in /tmp/dpng
    fp=os.path.join(ROOT,"diagrams",base+".png")
    if not os.path.exists(fp):
        fp=os.path.join(PNG, base+".png")
    if os.path.exists(fp):
        # downsample so the .docx stays emailable (owner PNGs are 1-3MB each at full res)
        im=Image.open(fp).convert("RGB"); maxw=1400
        if im.width>maxw: im=im.resize((maxw, round(im.height*maxw/im.width)), Image.LANCZOS)
        buf=io.BytesIO(); im.save(buf, format="JPEG", quality=85); buf.seek(0)
        doc.add_picture(buf, width=Inches(4.4))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False

def placeholder(line):
    s=line.strip().strip("`").strip()
    m=re.match(r'^\[(CATEGORY|ERA BAND|HERO PLATE|VISUAL|LAYOUT)\s*:?\s*(.*)\]$', s, re.S)
    if not m: return False
    kind,val=m.group(1),m.group(2).strip()
    if kind=="CATEGORY":
        p=doc.add_paragraph(); r=p.add_run(val.upper()); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=GRY
    elif kind=="ERA BAND":
        p=doc.add_paragraph(); r=p.add_run("Era:  "+val); r.italic=True; r.font.size=Pt(9)
    elif kind=="HERO PLATE":
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("[ HERO PHOTO TO COME — "+val+" ]"); r.italic=True; r.font.color.rgb=GRY; r.font.size=Pt(10)
    elif kind=="VISUAL":
        mref=re.search(r'([\w/\-.]+)\.(svg|md)', val)
        if not (mref and add_png(os.path.basename(mref.group(1)))):
            p=doc.add_paragraph(); r=p.add_run("[ Decision tree — see the diagram file ]"); r.italic=True; r.font.color.rgb=GRY
    elif kind=="LAYOUT":
        p=doc.add_paragraph(); r=p.add_run(val); r.italic=True; r.font.color.rgb=GRY
    return True

def render(md):
    lines=md.split("\n"); i=0
    while i<len(lines):
        line=lines[i]; s=line.strip()
        if placeholder(line): i+=1; continue
        if s=="": i+=1; continue
        if s.startswith("```"):
            i+=1
            while i<len(lines) and not lines[i].strip().startswith("```"): i+=1
            i+=1
            tree=FILE_TREE.get(CURRENT_FILE)
            if not (tree and add_png(tree)):
                p=doc.add_paragraph(); r=p.add_run("[ Decision tree — see the diagram file ]"); r.italic=True; r.font.color.rgb=GRY
            continue
        if re.match(r'^#{1,6}\s', s):
            lvl=len(s)-len(s.lstrip("#")); txt=s[lvl:].strip()
            doc.add_heading(txt, level=min(lvl,3)); i+=1; continue
        if s in ("---","***","___"): i+=1; continue
        if s.startswith(">"):
            buf=[]
            while i<len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].strip()); i+=1
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
            r=p.add_run(" ".join(buf)); r.italic=True; continue
        if "|" in s and i+1<len(lines) and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i+1]) and "-" in lines[i+1]:
            header=[c.strip() for c in s.strip().strip("|").split("|")]; i+=2; rows=[]
            while i<len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i+=1
            t=doc.add_table(rows=1, cols=len(header)); t.style='Light Grid Accent 1'
            for j,c in enumerate(header):
                cell=t.rows[0].cells[j]; cell.paragraphs[0].text=""
                run=cell.paragraphs[0].add_run(c); run.bold=True; run.font.size=Pt(9)
            for r in rows:
                cells=t.add_row().cells
                for j,c in enumerate(r):
                    if j<len(cells):
                        cells[j].paragraphs[0].text=""; add_runs(cells[j].paragraphs[0], c)
                        for rr in cells[j].paragraphs[0].runs: rr.font.size=Pt(9)
            continue
        if re.match(r'^[-*]\s', s):
            while i<len(lines) and re.match(r'^[-*]\s', lines[i].strip()):
                p=doc.add_paragraph(style='List Bullet'); add_runs(p, lines[i].strip()[2:]); i+=1
            continue
        if re.match(r'^\d+\.\s', s):
            while i<len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                p=doc.add_paragraph(style='List Number'); add_runs(p, re.sub(r'^\d+\.\s','',lines[i].strip())); i+=1
            continue
        def cont(ln):
            t=ln.strip()
            if not t: return False
            if re.match(r'^(#{1,6}\s|>|[-*]\s|\d+\.\s|```|\|)', t): return False
            if re.match(r'^`?\[(CATEGORY|ERA BAND|HERO PLATE|VISUAL|LAYOUT)', t): return False
            return True
        buf=[line]; i+=1
        while i<len(lines) and cont(lines[i]):
            buf.append(lines[i]); i+=1
        p=doc.add_paragraph(); add_runs(p, " ".join(b.strip() for b in buf))

# cover
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("Field Guide to American Furniture Identification, 1840–1940"); r.bold=True; r.font.size=Pt(22)
sb=doc.add_paragraph(); sb.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sb.add_run("Spotting, Dating & Valuing the Antiques You'll Actually Find — Victorian to Art Deco"); r.italic=True; r.font.size=Pt(13)
doc.add_page_break()

for title, files in SECTIONS:
    doc.add_page_break()
    h=doc.add_paragraph(); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=h.add_run(title); r.bold=True; r.font.size=Pt(20)
    for f in files:
        md=read(f)
        if not md: continue
        CURRENT_FILE=f
        doc.add_page_break()
        render(strip_chapter(md))
        for base in CHAPTER_PLATES.get(f, []):
            add_png(base)

doc.save(OUT)
print("wrote", OUT)
